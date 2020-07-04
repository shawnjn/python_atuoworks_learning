#!/usr/bin/env python
#-- coding: utf-8 --
#Time : 2020/7/5
#Author : shawn_jn

#常规工作，支付签字表格自动生成。--支付小助手v1.0.py

from docx import Document
from docx.enum.text import WD_LINE_SPACING, WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Pt, RGBColor #榜数、颜色
from docx.oxml.ns import qn #对应中文格式
import openpyxl as xl
import time

print('------注意：请再次确认是否正确填写《项目基本信息表》！------')
chose = 1
while chose == 1:
    print('【1】--已确认数据，继续执行！\n【2】--不执行，退出程序！')
    chose_in = input('请选择【1】or【2】：')
    print('\n')
    if chose_in == "1":
        ch = 2
        chose = 2
    elif chose_in == "2":
        ch = 1
        chose = 2
    else:
        print(f'"{chose_in}" 输入错误，请重新选择【1】or【2】\n')
if ch == 1:
    print('程序已退出，谢谢使用！')
if ch == 2:
    #支付审批表编辑，复制打开模板填写信息并另存为：合同名称-支付审批表(日期xxx）.xlsx文件
    today = time.strftime('%Y{y}%m{m}%d{d}', time.localtime()).format(y='年', m='月', d='日')

    wb1 = xl.load_workbook('./data_base/【不许修改】-支付审批表-官方2020年OA文件.xlsx')   #打开最新的"项目支付信息表"
    sheetw = wb1.worksheets[0]
    sheetw_copy = wb1.copy_worksheet(wb1.worksheets[0])
    sheetw_copy.title = today
    # wb1.remove_sheet(wb1.get_sheet_by_name('sheet'))    #删除多余sheet
    del wb1['sheet']    #删除多余sheet

    #打开项目支付信息表
    wb2 = xl.load_workbook('【填写】-项目基本信息表-V1.0.xlsx', data_only=True)  #data_only=True表示只获取数据，默认Fasle
    wb2_data = wb2.worksheets[0]

    #在sheetw_copy中输入数据
    sheetw_copy.cell(4, 4).value = today    #获取当前时间（填表时间）
    sheetw_copy.cell(6, 3).value = wb2_data.cell(1, 3).value    #获取合同/规费名称
    sheetw_copy.cell(7, 3).value = wb2_data.cell(2, 3).value    #获取合同编号
    sheetw_copy.cell(7, 7).value = wb2_data.cell(5, 3).value    #获取合同金额
    sheetw_copy.cell(8, 3).value = wb2_data.cell(3, 3).value    #获取收款单位
    sheetw_copy.cell(8, 6).value = wb2_data.cell(4, 3).value    #获取开户行信息
    sheetw_copy.cell(9, 3).value = wb2_data.cell(8, 7).value    #获取累计已支付预付款
    sheetw_copy.cell(9, 6).value = wb2_data.cell(6, 7).value    #获取本期申请支付首付款
    sheetw_copy.cell(10, 3).value = wb2_data.cell(6, 3).value    #获取累计已结算金额
    sheetw_copy.cell(10, 6).value = wb2_data.cell(7, 3).value    #获取本次结算金额
    sheetw_copy.cell(12, 3).value = wb2_data.cell(1, 7).value    #获取预付款
    sheetw_copy.cell(12, 5).value = wb2_data.cell(2, 7).value    #获取投运款
    sheetw_copy.cell(12, 6).value = wb2_data.cell(3, 7).value    #获取质保金
    sheetw_copy.cell(12, 7).value = wb2_data.cell(4, 7).value    #获取考核或其他金额
    sheetw_copy.cell(13, 8).value = wb2_data.cell(10, 3).value    #获取本次应支付金额

    wb1.save(f'《{wb2_data.cell(1, 3).value }》-支付审批表{today}.xlsx') #保存支付审批表
    wb1.close()
    print(f'------{wb2_data.cell(1, 3).value }支付审批表已填写完成------')

    #设置真整个文档的字体和段落基本格式
    document = Document()
    document.styles['Normal'].font.name = u'方正仿宋_GBK'
    document.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), u'方正仿宋_GBK')
    # today1 = time.strftime('%Y-%m-%d', time.localtime())
    #由于time库对中文支持不好，可以用以下方式处理
    today = time.strftime('%Y{y}%m{m}%d{d}', time.localtime()).format(y='年', m='月', d='日')
    # contract_name = input('请输入合同名称：')   #后续从数据库内筛选

    #编辑word内容
    p = document.add_paragraph()
    run_01 = p.add_run('xxxxxx有限公司\n合同款付款说明')
    run_01.font.name = u'方正小标宋_GBK'
    run_01.element.rPr.rFonts.set(qn('w:eastAsia'), '方正小标宋_GBK')
    run_01.font.size = Pt(22)  # 字体大小
    p.paragraph_format.line_spacing = Pt(28) #行间距设置

    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    #创建表格,并将表格在word页面中居中
    table = document.add_table(rows=5, cols=4, style="TableGrid")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # table.style.font.name = u'方正仿宋_GBK'
    textlist = [
        '合同/征收单位', '合同/规费名称', '合同/规费总价', '合同编号', '上期累计支付', '本期申请支付', '本期累计支付', '本期支付说明']
    n = 0
    for i in range(4):
        for j in (0, 2):
            table.cell(i, j).text = textlist[n]
            n += 1
            # print(table.cell(i, j).text, n)
    table.cell(0, 1).text = wb2_data.cell(3, 3).value   #收款单位
    table.cell(0, 3).text = wb2_data.cell(1, 3).value   #合同名称
    table.cell(1, 1).text = str('%.2f' % wb2_data.cell(5, 3).value) + '元'  #合同总价
    table.cell(1, 3).text = str(wb2_data.cell(2, 3).value)  #合同编号
    table.cell(2, 1).text = str('%.2f' % wb2_data.cell(6, 3).value) + '元'   #上期累计支付
    table.cell(2, 3).text = str('%.2f' % wb2_data.cell(10, 3).value) + '元'   #本期申请支付
    table.cell(3, 1).text = str('%.2f' % wb2_data.cell(11, 3).value) + '元'   #本期累计支付
    table.cell(3, 3).text = wb2_data.cell(9, 3).value   #本期支付说明
    # 表格排版格式进行调整，left格式
    for row in range(0, 4):
        for col in range(0, 4):
            table.cell(row, col).vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraphs_table = table.cell(row, col).paragraphs[0]
            paragraphs_table.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT    #等同paragraphs_table.alignment = 0
                                                                    # 也可（for left, 1 for right, 2 center, 3 justify ....）


    table.cell(4, 0).merge(table.cell(4, 3))    #合并单元格
    table_parag_1 = table.cell(4, 0).paragraphs[0] #将合并单元
    table_parag_1.paragraph_format.line_spacing =Pt(20)
    #(2)数据输入，最后格表格输入

    run_1_1 = table_parag_1.add_run('    1.支付条款，合同第')
    run_1_2 = table_parag_1.add_run('%s' % wb2_data.cell(2, 10).value).font.underline = True    #第X条
    run_1_3 = table_parag_1.add_run('条第')
    run_1_4 = table_parag_1.add_run('%s' % wb2_data.cell(2, 12).value).font.underline = True    #第X款
    run_1_5 = table_parag_1.add_run('款约定:')
    num = int(wb2_data.cell(2, 12).value)
    text_t =wb2_data.cell(num+1, 14).value

    run_1_6 = table_parag_1.add_run(text_t).font.underline = True   #合同支付条款

    table_parag_2 = table.cell(4, 0).add_paragraph()
    run_2_1 = table_parag_2.add_run('    2.')
    run_2_2 = table_parag_2.add_run(today).font.underline = True
    run_2_3 = table_parag_2.add_run(',乙方')
    run_2_4 = table_parag_2.add_run(wb2_data.cell(9, 3).value).font.underline = True    #验收xxx事项，说明
    run_2_5 = table_parag_2.add_run('事项通过了我公司')
    run_2_6 = table_parag_2.add_run('运维检修部').font.underline = True
    run_2_7 = table_parag_2.add_run('验收，符合合同约定的结算支付条件。')

    table_parag_3 = table.cell(4, 0).add_paragraph()
    run_3_1 = table_parag_3.add_run('    3.乙方已提交相应的')
    text_elements =[]
    for i in range(5):
        if wb2_data.cell(i+7, 14).value != None:
            text_elements.append(wb2_data.cell(i+7, 14).value)
    text1 =''
    text2 =''
    text3 =''
    i = 1
    for t in range(len(text_elements)-1):
        text1 = text_elements[t] + '、' + text1
        text3 = text3 + str(i) + '.' + text_elements[t] +'、'
        i += 1
    text2 = text1 + text_elements[-1]
    text3 = text3 + str(i) +'.' + text_elements[-1]
    # text1 = input('付款申请、发票及xxx资料：')
    run_3_2 = table_parag_3.add_run('%s' % text2).font.underline = True
    run_3_3 = table_parag_3.add_run('等，符合合同约定。')

    table_parag_4 = table.cell(4, 0).add_paragraph()
    run4_1 = table_parag_4.add_run('    4. 本次结算金额为：')
    run4_2 = table_parag_4.add_run(str('%.2f' % wb2_data.cell(7, 3).value)).font.underline = True #输入结算金额
    run4_3 = table_parag_4.add_run('元(不含税金额为：')
    run4_4 = table_parag_4.add_run(str('%.2f' % (wb2_data.cell(7, 3).value/(1+wb2_data.cell(8, 3).value)))) #不含税金额
    run4_4.font.underline = True
    run4_5 = table_parag_4.add_run('元），')
    run4_4 = table_parag_4.add_run('应扣款项为：')
    sum_decrees = 0
    for m in range(4):
        sum_decrees += wb2_data.cell(m+1, 7).value
    run4_5 = table_parag_4.add_run(str('%.2f' % sum_decrees) + '元').font.underline = True
    run4_6 = table_parag_4.add_run('，应支付：')
    run4_7 = table_parag_4.add_run(str('%.2f' % wb2_data.cell(10, 3).value) + '元').font.underline = True    #应支付金额
    run4_8 = table_parag_4.add_run('，累计支付金额为：')
    run4_9 = table_parag_4.add_run(str('%.2f' % wb2_data.cell(11, 3).value) + '元').font.underline = True    #累计支付金额

    table_parag_5 = table.cell(4, 0).add_paragraph()
    run5_1 = table_parag_5.add_run('    本期累计支付占合同总价的比例为：')
    per = wb2_data.cell(11, 3).value*100/wb2_data.cell(5, 3).value
    run5_2 = table_parag_5.add_run(str(round(per, 2)) + '%。') #支付比例
    run5_2.font.underline = True

    table_parag_6 = table.cell(4, 0).add_paragraph()
    run6_1 = table_parag_6.add_run('    附  件： （合同规定的有关材料）\n')

    table_parag_7 = table.cell(4, 0).add_paragraph('')
    run7_1 = table_parag_6.add_run('    ' + text3)

    table_parag_8 = table.cell(4, 0).add_paragraph('\n\n\n经办人：\n审核：\n运维检修部')
    table_parag_8.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    table_parag_8.paragraph_format.right_indent = Pt(60)

    table_parag_9 = table.cell(4, 0).add_paragraph('年  月  日\n\n')
    table_parag_9.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    table_parag_9.paragraph_format.right_indent = Pt(30)

    document.save(f'%s-{today}-付款说明.docx' % wb2_data.cell(1, 3).value)
    print(f'------{wb2_data.cell(1, 3).value}付款说明已编制完成！！！------')
    wb2.close()

